from docx import Document
import os
import uuid
from num2words import num2words

# Conversión de números a letras
def numero_a_letras(cantidad):
    return num2words(cantidad, lang='es').replace("-", " ").replace(",", "").replace(".", "").upper()

# Reemplazo de etiquetas en párrafos
def reemplazar_etiquetas_en_parrafos(parrafos, reemplazos):
    for p in parrafos:
        for run in p.runs:
            for key, val in reemplazos.items():
                if key in run.text:
                    run.text = run.text.replace(key, str(val))

# Buscar tabla de productos
def obtener_tabla_productos(tablas, columnas, tipo_formulario):
    for table in tablas:
        try:
            if len(table.columns) == columnas:
                if tipo_formulario == "andy":
                    for row in table.rows:
                        for cell in row.cells:
                            if "{cantidad_1}" in cell.text or "{descripcion_1}" in cell.text:
                                return table
                else:  # ecualimpio u otros
                    return table
        except Exception:
            continue
    return None

# Reemplazo de etiquetas en tabla de productos
def reemplazar_etiquetas_en_tabla_productos(productos, table, tipo_formulario):
    if not productos:
        return

    if tipo_formulario == "andy":
        tags = {
            "{cantidad_1}": str(productos[0]["cantidad"]),
            "{cpc_1}": productos[0]["cpc"],
            "{descripcion_1}": productos[0]["descripcion"],
            "{precio_unitario_1}": f"${float(productos[0]['precio_unitario']):.2f}",
            "{total_1}": f"${float(productos[0]['total']):.2f}",
        }
    elif tipo_formulario == "ecualimpio":
        tags = {
            "{productos_1}": 1,
            "{cpc_1}": productos[0]["cpc"],
            "{descripcion_1}": productos[0]["descripcion"],
            "{cantidad_1}": str(productos[0]["cantidad"]),
            "{precio_unitario_1}": f"${float(productos[0]['precio_unitario']):.2f}",
            "{total_1}": f"${float(productos[0]['total']):.2f}",
        }
    else:
        raise ValueError(f"Tipo de formulario '{tipo_formulario}' no soportado.")

    for row in table.rows:
        for cell in row.cells:
            for key, val in tags.items():
                if key in cell.text:
                    cell.text = cell.text.replace(key, str(val))

# Agregar más filas para los productos extra
def agregar_filas_para_productos(productos, table, tipo_formulario):
    if not productos:
        return

    base_row_idx = 1 if tipo_formulario == "andy" else 2
    base_row = table.rows[base_row_idx]

    for i, prod in enumerate(productos):
        if i == 0:
            continue  # ya está el primer producto en la plantilla

        row = table.add_row()
        for j in range(len(base_row.cells)):
            row.cells[j].text = base_row.cells[j].text

        if tipo_formulario == "andy":
            row.cells[0].text = str(i + 1)
            row.cells[1].text = str(prod["cantidad"])
            row.cells[2].text = prod["cpc"]
            row.cells[3].text = prod["descripcion"]
            row.cells[4].text = f"${float(prod['total']):.2f}"
        else:  # ecualimpio
            row.cells[0].text = str(i + 1)
            row.cells[1].text = str(prod["cpc"])
            row.cells[2].text = prod["descripcion"]
            row.cells[3].text = str(prod["cantidad"])
            row.cells[4].text = f"${float(prod['precio_unitario']):.2f}"
            row.cells[5].text = f"${float(prod['total']):.2f}"

# Reemplazo de etiquetas generales en otras tablas
def reemplazar_etiquetas_generales_en_tablas(tablas, reemplazos):
    for table in tablas:
        if len(table.columns) in [5, 6] and len(table.rows) >= 1:
            continue  # skip tabla de productos
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for key, val in reemplazos.items():
                            if key in run.text:
                                run.text = run.text.replace(key, str(val))

# Cálculo de totales
def calcular_totales(data):
    productos = data["productos"]
    subtotal = sum(float(p["total"]) for p in productos)
    iva = float(data.get("iva", 0))
    iva_valor = round((subtotal * iva) / 100, 2)
    total = round(subtotal + iva_valor, 2)

    data["subtotal"] = f"${subtotal:.2f}"
    data["iva"] = f"${iva_valor:.2f}"
    data["total"] = f"${total:.2f}"
    data["letras"] = numero_a_letras(subtotal)

# Generador principal de proforma
def generar_proforma_docx(data, plantilla_path):
    doc = Document(plantilla_path)

    calcular_totales(data)

    reemplazar_etiquetas_en_parrafos(doc.paragraphs, data)

    tipo = data.get("tipo_formulario")
    columnas = 5 if tipo == "andy" else 6

    table = obtener_tabla_productos(doc.tables, columnas, tipo)
    if not table:
        raise Exception("No se encontró tabla de productos adecuada.")

    reemplazar_etiquetas_en_tabla_productos(data["productos"], table, tipo)

    if len(data["productos"]) > 1:
        agregar_filas_para_productos(data["productos"], table, tipo)

    reemplazar_etiquetas_generales_en_tablas(doc.tables, data)

    file_id = uuid.uuid4().hex[:8]
    output_filename = f"proforma_{file_id}.docx"
    output_path = os.path.join("/descargas", output_filename)
    doc.save(output_path)

    return output_path